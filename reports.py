import pandas as pd
from docx import Document
from docx.shared import Inches
import os


def generate_word_report(report_data, mission_info, output_path):
    doc = Document()
    logo = "static/logo.png"
    if os.path.exists(logo): doc.add_image(logo, width=Inches(1.8))

    doc.add_heading('Neosky India Ltd - Incident Analysis', 0)
    doc.add_heading('1.0 Mission Details', level=1)
    table = doc.add_table(rows=5, cols=2)
    info = [
        ("Drone Model:", mission_info.get('drone')),
        ("MSN Number:", mission_info.get('msn')),
        ("Incident Date:", mission_info.get('date')),
        ("File Name:", report_data['filename']),
        ("Health Status:", report_data['status'])
    ]
    for i, (k, v) in enumerate(info):
        table.rows[i].cells[0].text, table.rows[i].cells[1].text = k, str(v)

    if report_data['alerts']:
        doc.add_heading('CRITICAL ALERTS DETECTED', level=1)
        for alert in report_data['alerts']:
            p = doc.add_paragraph(f" {alert}")
            p.runs[0].bold = True

    current_cat = ""
    for item in report_data['details']:
        if item['cat'] != current_cat:
            current_cat = item['cat']
            doc.add_heading(current_cat, level=1)
            t = doc.add_table(rows=1, cols=5)
            t.style = 'Table Grid'
            for i, txt in enumerate(['Parameter', 'Min', 'Max', 'Avg', 'Dev%']): t.rows[0].cells[i].text = txt

        row = t.add_row().cells
        row[0].text, row[1].text, row[2].text, row[3].text, row[4].text = item['param'], str(item['min']), str(
            item['max']), str(item['avg']), f"{item['dev']}%"

    doc.add_page_break()
    doc.add_heading('Appendix: Full Parameter Dump', level=1)
    p_table = doc.add_table(rows=1, cols=2)
    p_table.style = 'Table Grid'
    for name, val in report_data['params'].items():
        r = p_table.add_row().cells
        r[0].text, r[1].text = str(name), str(val)

    doc.save(output_path)

def generate_excel(report_data, output_path):
    with pd.ExcelWriter(output_path) as writer:
        pd.DataFrame(report_data['details']).to_excel(writer, sheet_name='Telemetry', index=False)
        pd.DataFrame(list(report_data['params'].items())).to_excel(writer, sheet_name='Parameters', index=False)
        pd.DataFrame(report_data['events'], columns=['Messages']).to_excel(writer, sheet_name='Messages', index=False)